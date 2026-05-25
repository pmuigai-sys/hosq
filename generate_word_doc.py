import os
import glob
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx not found. Please install it first.")
    exit(1)

def add_markdown_to_doc(doc, md_file):
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    filename = os.path.basename(md_file).replace('.md', '').replace('-', ' ').title()
    doc.add_heading(filename, level=1)
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('###'):
            doc.add_heading(line.replace('###', '').strip(), level=3)
        elif line.startswith('##'):
            doc.add_heading(line.replace('##', '').strip(), level=2)
        elif line.startswith('#'):
            continue # already added title
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. '):
            doc.add_paragraph(line[3:], style='List Number')
        else:
            doc.add_paragraph(line)
    
    doc.add_page_break()

def generate_master_doc():
    doc = Document()
    
    # Professional Styling
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Title Page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\n\n\n\n\n\n\n\nHOSQ: Hospital Queue System\n')
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(41, 128, 185) # Professional Blue
    run.bold = True
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('Official Project Documentation and Technical Manual\n')
    run2.font.size = Pt(18)
    
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(f'Author: Peter Thairu Muigai\nDate: 2026-02-18')
    run3.font.size = Pt(12)
    
    doc.add_page_break()
    
    # Table of Contents placeholder (Manual in Word usually)
    doc.add_heading('Table of Contents', level=1)
    doc.add_paragraph('1. System Architecture')
    doc.add_paragraph('2. Integration Details')
    doc.add_paragraph('3. Setup Guide')
    doc.add_paragraph('4. SMS Provider Comparison')
    doc.add_paragraph('5. Workflow Guidelines')
    doc.add_page_break()
    
    # Add files in order
    files_to_add = [
        'f:/hosq/docs/architecture.md',
        'f:/hosq/docs/integrations.md',
        'f:/hosq/SETUP_GUIDE.md',
        'f:/hosq/docs/sms-alternatives.md',
        'f:/hosq/docs/workflows.md'
    ]
    
    for file_path in files_to_add:
        if os.path.exists(file_path):
            add_markdown_to_doc(doc, file_path)
            
    output = 'f:/hosq/HOSQ_Project_Full_Documentation.docx'
    doc.save(output)
    print(f"✅ Master Document Generated: {output}")

if __name__ == "__main__":
    generate_master_doc()
